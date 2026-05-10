from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# --- Font mặc định ---
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)

def set_font(run, bold=False, size=13, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    h = doc.add_heading(level=level)
    h.clear()
    run = h.add_run(text)
    sizes = {1: 16, 2: 14, 3: 13}
    set_font(run, bold=True, size=sizes.get(level, 13))
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(20)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=12)
    doc.add_paragraph()

# ===================== NỘI DUNG =====================

# Tiêu đề chính
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('BÁO CÁO TỔNG QUAN DỰ ÁN')
set_font(r, bold=True, size=18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('CAMPUS MARKETPLACE ĐÀ LẠT')
set_font(r2, bold=True, size=16)

doc.add_paragraph()

# Thông tin chung
info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Table Grid'
infos = [
    ('Tên dự án', 'Campus Marketplace Đà Lạt'),
    ('Tên kỹ thuật', 'campus-marketplace'),
    ('Phiên bản', '0.1.0'),
    ('Ngày lập báo cáo', '10/05/2026'),
    ('Loại dự án', 'Dự án cá nhân (Personal Project)'),
]
for i, (k, v) in enumerate(infos):
    c1 = info_table.rows[i].cells[0]
    c2 = info_table.rows[i].cells[1]
    c1.text = ''
    c2.text = ''
    set_font(c1.paragraphs[0].add_run(k), bold=True, size=12)
    set_font(c2.paragraphs[0].add_run(v), size=12)
doc.add_paragraph()

# 1. TỔNG QUAN
add_heading('1. TỔNG QUAN DỰ ÁN', 1)
add_para(
    'Campus Marketplace Đà Lạt là một nền tảng chợ trực tuyến dành riêng cho sinh viên tại khu vực Đà Lạt và Bảo Lộc. '
    'Ứng dụng cho phép sinh viên mua, bán và trao đổi đồ dùng học tập, thiết bị, tài liệu — đặc biệt với tính năng '
    'tìm kiếm theo mã học phần (course code), giúp khớp chính xác nhu cầu theo từng môn học.',
    indent=True
)
add_para('Điểm khác biệt nổi bật so với các sàn thương mại điện tử thông thường:', bold=True)
add_bullet('Chỉ dành riêng cho sinh viên có email đuôi .edu.vn')
add_bullet('Tìm kiếm và lọc theo mã học phần, chuyên ngành, địa điểm')
add_bullet('Tập trung vào cộng đồng sinh viên địa phương (Đà Lạt, Bảo Lộc)')

# 2. MỤC TIÊU
add_heading('2. MỤC TIÊU DỰ ÁN', 1)
add_heading('2.1. Mục Tiêu Chính', 2)
add_table(
    ['STT', 'Mục tiêu', 'Mô tả'],
    [
        ('1', 'Tạo sân chơi trao đổi đồ học tập', 'Nơi sinh viên có thể mua/bán sách giáo trình, thiết bị thực hành, tài liệu học tập cũ'),
        ('2', 'Tìm kiếm theo mã học phần', 'Cho phép tìm đúng sản phẩm liên quan đến môn học (IT203, CH102, EC101...)'),
        ('3', 'Xác thực danh tính sinh viên', 'Chỉ email .edu.vn mới đăng tin được, tăng độ tin cậy cộng đồng'),
        ('4', 'Hỗ trợ đăng tin nhiều ảnh', 'Upload nhiều ảnh thực tế lên Supabase Storage'),
        ('5', 'Tối ưu SEO địa phương', 'Metadata động tối ưu cho từng từ khoá tìm kiếm'),
    ]
)
add_heading('2.2. Mục Tiêu Kỹ Thuật', 2)
add_bullet('Xây dựng web app hiện đại với Next.js 14 App Router và TypeScript')
add_bullet('Sử dụng Supabase làm backend tích hợp (Auth + Database PostgreSQL + Storage)')
add_bullet('Xây dựng UI nhanh với Tailwind CSS và shadcn/ui style')
add_bullet('Hệ thống fallback data: khi chưa kết nối Supabase, app vẫn hiển thị 50 sản phẩm mẫu')
add_bullet('Bảo vệ route nhạy cảm bằng Next.js Middleware')

# 3. KIẾN TRÚC
add_heading('3. KIẾN TRÚC KỸ THUẬT', 1)
add_heading('3.1. Technology Stack', 2)
add_table(
    ['Lớp', 'Công nghệ', 'Phiên bản'],
    [
        ('Frontend Framework', 'Next.js (App Router)', '14.2.5'),
        ('Ngôn ngữ', 'TypeScript', '^5.6.3'),
        ('UI Styling', 'Tailwind CSS', '^3.4.15'),
        ('Icon', 'Lucide React', '^0.451.0'),
        ('Backend / Auth', 'Supabase', '^2.45.0'),
        ('Auth Helpers', '@supabase/auth-helpers-nextjs', '^0.9.0'),
        ('Utility', 'clsx, tailwind-merge, class-variance-authority', '—'),
    ]
)
add_heading('3.2. Kiến Trúc Tổng Thể', 2)
add_para('Ứng dụng được xây dựng theo mô hình phân lớp rõ ràng:', indent=True)
add_bullet('Tầng giao diện (Frontend): Next.js App Router với các trang chính: Trang chủ (/), Tìm kiếm (/search), Chi tiết sản phẩm (/products/[id]), Đăng tin (/sell).')
add_bullet('Tầng xác thực (Middleware): Kiểm tra domain email .edu.vn và bảo vệ route yêu cầu đăng nhập.')
add_bullet('Tầng backend (Supabase): Xác thực OTP qua email, cơ sở dữ liệu PostgreSQL, lưu trữ ảnh (Storage).')

# 4. CƠ SỞ DỮ LIỆU
add_heading('4. CƠ SỞ DỮ LIỆU (DATABASE SCHEMA)', 1)
add_heading('4.1. Bảng profiles — Thông Tin Người Dùng', 2)
add_para('Lưu thông tin người dùng, được tạo tự động khi đăng ký qua trigger.', indent=True)
add_table(
    ['Cột', 'Kiểu dữ liệu', 'Mô tả'],
    [
        ('id', 'UUID (PK)', 'Liên kết với auth.users'),
        ('full_name', 'text', 'Tên đầy đủ'),
        ('avatar_url', 'text', 'URL ảnh đại diện'),
        ('university_city', 'text', 'Thành phố (mặc định: Đà Lạt)'),
        ('major', 'text', 'Chuyên ngành'),
        ('created_at', 'timestamptz', 'Thời gian tạo'),
        ('updated_at', 'timestamptz', 'Thời gian cập nhật (tự động)'),
    ]
)
add_heading('4.2. Bảng products — Sản Phẩm / Tin Đăng', 2)
add_table(
    ['Cột', 'Kiểu dữ liệu', 'Mô tả'],
    [
        ('id', 'UUID (PK)', 'ID sản phẩm'),
        ('user_id', 'UUID (FK)', 'Người đăng (→ profiles.id)'),
        ('title', 'text', 'Tên sản phẩm'),
        ('price', 'numeric(12,2)', 'Giá (VND, ≥ 0)'),
        ('category', 'text', 'Danh mục (CNTT, Hóa học, Kinh tế...)'),
        ('status', 'text', 'Trạng thái: available / reserved / sold'),
        ('university_city', 'text', 'Địa điểm (Đà Lạt, Bảo Lộc)'),
        ('course_code', 'text', 'Mã học phần (VD: IT203, CH102)'),
        ('description', 'text', 'Mô tả chi tiết'),
        ('image_urls', 'text[]', 'Mảng URL ảnh từ Supabase Storage'),
        ('created_at', 'timestamptz', 'Ngày đăng'),
    ]
)
add_para('Indexes: course_code, university_city, category, status — tối ưu hiệu năng tìm kiếm và lọc.')

add_heading('4.3. Bảng messages — Tin Nhắn', 2)
add_para('Lưu tin nhắn giữa người mua và người bán (đã thiết kế schema, chưa có UI).', indent=True)
add_table(
    ['Cột', 'Kiểu dữ liệu', 'Mô tả'],
    [
        ('id', 'UUID (PK)', 'ID tin nhắn'),
        ('product_id', 'UUID (FK)', 'Sản phẩm liên quan'),
        ('sender_id', 'UUID (FK)', 'Người gửi'),
        ('receiver_id', 'UUID (FK)', 'Người nhận'),
        ('content', 'text', 'Nội dung tin nhắn'),
        ('read_at', 'timestamptz', 'Thời điểm đã đọc (nullable)'),
        ('created_at', 'timestamptz', 'Thời điểm gửi'),
    ]
)
add_heading('4.4. Triggers & Functions', 2)
add_table(
    ['Tên', 'Loại', 'Mô tả'],
    [
        ('handle_new_user()', 'Function + Trigger', 'Tự động tạo bản ghi profiles khi user đăng ký'),
        ('set_updated_at()', 'Function + Trigger', 'Tự động cập nhật updated_at khi có thay đổi'),
    ]
)

# 5. CHỨC NĂNG
add_heading('5. PHÂN TÍCH CHỨC NĂNG', 1)
add_heading('5.1. FC-01: Xác Thực & Phân Quyền', 2)
add_para('Hệ thống xác thực sinh viên qua email .edu.vn, bảo vệ các route cần đăng nhập.', indent=True)
add_para('Luồng xử lý:', bold=True)
add_bullet('Người dùng truy cập bất kỳ route nào → Middleware lấy session Supabase')
add_bullet('Nếu có session → kiểm tra email có đuôi .edu.vn; không hợp lệ → redirect /auth/blocked')
add_bullet('Nếu truy cập /sell mà chưa đăng nhập → redirect /auth/login?next=/sell')

add_heading('5.2. FC-02: Trang Chủ (Home)', 2)
add_para('Trang landing chính, tổng hợp tất cả tính năng trọng tâm của ứng dụng.', indent=True)
add_table(
    ['Query Param', 'Mô tả'],
    [
        ('?q=', 'Tìm kiếm theo từ khoá'),
        ('?major=', 'Lọc theo chuyên ngành'),
        ('?city=', 'Lọc theo địa điểm'),
        ('?course=', 'Lọc theo mã học phần'),
    ]
)

add_heading('5.3. FC-03: Tìm Kiếm Nâng Cao', 2)
add_para('Trang tìm kiếm chuyên biệt, tối ưu cho mã học phần, có SEO metadata động.', indent=True)
add_bullet('Form tìm kiếm theo mã học phần (?q=IT203)')
add_bullet('Dynamic metadata: title và description thay đổi theo từ khoá')
add_bullet('Ví dụ title: "Tìm kiếm IT203 | Campus Marketplace Đà Lạt"')

add_heading('5.4. FC-04: Danh Sách & Chi Tiết Sản Phẩm', 2)
add_para('Hiển thị danh sách dạng lưới và trang chi tiết từng sản phẩm.', indent=True)
add_bullet('ProductCard: ảnh bìa, badge chuyên ngành, badge trạng thái, tên, giá (VND), địa điểm, mã học phần')
add_bullet('Trang /products/[id]: ảnh full-width, giá lớn, mô tả đầy đủ, gợi ý liên hệ người bán')
add_bullet('Logic: có Supabase → query DB; không có → dùng 50 sản phẩm mẫu (fallback)')

add_heading('5.5. FC-05: Bộ Lọc Sản Phẩm', 2)
add_table(
    ['Bộ lọc', 'Kiểu', 'Giá trị'],
    [
        ('Chuyên ngành', 'Dropdown', 'CNTT, Hóa học, Kinh tế, Sinh học, Khác'),
        ('Địa điểm', 'Dropdown', 'Đà Lạt, Bảo Lộc, Khác'),
        ('Mã học phần', 'Text input', 'VD: IT203, CH102...'),
        ('Từ khoá', 'Text input', 'Tìm trong tiêu đề + mã học phần'),
    ]
)

add_heading('5.6. FC-06: Đăng Tin Sản Phẩm (Yêu cầu đăng nhập)', 2)
add_table(
    ['Trường', 'Bắt buộc', 'Mô tả'],
    [
        ('Tiêu đề', 'Có', 'Tên sản phẩm'),
        ('Giá (VND)', 'Có', 'Giá bán'),
        ('Chuyên ngành', 'Có', 'Danh mục (dropdown)'),
        ('Địa điểm', 'Có', 'Đà Lạt / Bảo Lộc / Khác'),
        ('Mã học phần', 'Không', 'VD: IT101 (tùy chọn)'),
        ('Mô tả', 'Không', 'Mô tả chi tiết'),
        ('Ảnh sản phẩm', 'Không', 'Multiple file upload'),
    ]
)

add_heading('5.7. Trạng Thái Sản Phẩm', 2)
add_table(
    ['Trạng thái', 'Nhãn hiển thị', 'Ý nghĩa'],
    [
        ('available', 'Còn bán', 'Sản phẩm đang chờ người mua'),
        ('reserved', 'Đã giữ', 'Đang thương lượng / giữ chỗ'),
        ('sold', 'Đã bán', 'Giao dịch hoàn tất'),
    ]
)

# 6. LUỒNG NGƯỜI DÙNG
add_heading('6. LUỒNG NGƯỜI DÙNG (USER FLOWS)', 1)
add_heading('6.1. Luồng Người Mua', 2)
add_bullet('Vào trang chủ → Xem sản phẩm mới nhất')
add_bullet('Lọc theo chuyên ngành / địa điểm / mã học phần')
add_bullet('Click vào sản phẩm → Xem trang chi tiết')
add_bullet('Liên hệ người bán (tính năng chat đang trong kế hoạch)')

add_heading('6.2. Luồng Người Bán', 2)
add_bullet('Vào trang chủ → Click "Đăng tin miễn phí"')
add_bullet('Middleware: chưa đăng nhập → redirect /auth/login')
add_bullet('Đăng nhập OTP bằng email .edu.vn → Middleware kiểm tra domain → hợp lệ')
add_bullet('Redirect về /sell → Điền form + upload ảnh → Submit → lưu vào Supabase')
add_bullet('Sản phẩm hiển thị trên trang chủ')

add_heading('6.3. Luồng Email Không Hợp Lệ', 2)
add_bullet('Đăng nhập bằng email không có đuôi .edu.vn')
add_bullet('Middleware phát hiện domain sai → Redirect /auth/blocked?reason=domain')
add_bullet('Hiển thị thông báo tài khoản bị chặn')

# 7. CẤU HÌNH
add_heading('7. CẤU HÌNH & TRIỂN KHAI', 1)
add_heading('7.1. Biến Môi Trường', 2)
add_table(
    ['Biến', 'Mô tả'],
    [
        ('NEXT_PUBLIC_SUPABASE_URL', 'URL project Supabase'),
        ('NEXT_PUBLIC_SUPABASE_ANON_KEY', 'Anon key public của Supabase'),
    ]
)
add_heading('7.2. Thiết Lập Supabase', 2)
add_bullet('Database: Chạy schema.sql trong SQL Editor của Supabase')
add_bullet('Storage: Tạo bucket product-images và bật public access')
add_bullet('Auth: Bật xác thực OTP email')

add_heading('7.3. Các Scripts Thường Dùng', 2)
add_table(
    ['Lệnh', 'Mô tả'],
    [
        ('npm run dev', 'Chạy môi trường development (localhost:3000)'),
        ('npm run build', 'Build production bundle'),
        ('npm run start', 'Chạy server production'),
        ('npm run lint', 'Kiểm tra lỗi ESLint'),
        ('npm run seed', 'Seed dữ liệu mẫu vào Supabase'),
    ]
)

# 8. ĐIỂM MẠNH & HẠN CHẾ
add_heading('8. ĐIỂM MẠNH & HẠN CHẾ', 1)
add_heading('8.1. Điểm Mạnh', 2)
add_bullet('Fallback data thông minh: app hoạt động dù chưa cấu hình Supabase')
add_bullet('Bảo mật cộng đồng: chỉ email .edu.vn mới đăng tin được, giảm spam')
add_bullet('Tìm kiếm theo mã học phần: tính năng đặc thù, phù hợp nhu cầu sinh viên')
add_bullet('SEO địa phương: metadata động tối ưu cho từng từ khoá tìm kiếm')
add_bullet('TypeScript toàn bộ: type-safe từ database đến UI')
add_bullet('Server Components: tận dụng Next.js App Router để fetch data phía server')

add_heading('8.2. Hạn Chế & Chức Năng Chưa Hoàn Thiện', 2)
add_table(
    ['Tính năng', 'Trạng thái', 'Ghi chú'],
    [
        ('Chat / Nhắn tin', 'Chưa có UI', 'Schema messages đã tạo, chưa implement'),
        ('Quản lý tin đăng', 'Chưa có', 'Người đăng chưa thể sửa/xoá tin'),
        ('Trang Profile cá nhân', 'Chưa có', '—'),
        ('Phân trang (Pagination)', 'Chưa có', 'Hiện load toàn bộ sản phẩm'),
        ('Thông báo (Notification)', 'Chưa có', '—'),
        ('Bộ lọc theo giá', 'Chưa có', '—'),
        ('Lọc theo trạng thái', 'Chưa có', '—'),
        ('Xác minh email tự động', 'Một phần', 'Kiểm tra domain, chưa verify sinh viên'),
    ]
)

# 9. ROADMAP
add_heading('9. HƯỚNG PHÁT TRIỂN TIẾP THEO (ROADMAP)', 1)
add_heading('Giai Đoạn 1 — Hoàn Thiện Core', 2)
add_bullet('Implement trang quản lý tin đăng của người dùng (/my-listings)')
add_bullet('Thêm chức năng sửa / xoá tin đăng')
add_bullet('Thêm phân trang hoặc infinite scroll')
add_bullet('Bộ lọc theo khoảng giá')

add_heading('Giai Đoạn 2 — Tăng Tương Tác', 2)
add_bullet('Implement UI nhắn tin (sử dụng bảng messages đã có)')
add_bullet('Thông báo real-time khi có tin nhắn mới (Supabase Realtime)')
add_bullet('Trang profile cá nhân với lịch sử giao dịch')

add_heading('Giai Đoạn 3 — Mở Rộng', 2)
add_bullet('Hỗ trợ thêm nhiều thành phố (Nha Trang, Đà Nẵng...)')
add_bullet('Thêm hệ thống đánh giá người bán')
add_bullet('PWA (Progressive Web App) để trải nghiệm mobile tốt hơn')
add_bullet('Admin dashboard kiểm duyệt tin đăng')

# 10. GHI CHÚ
add_heading('10. GHI CHÚ', 1)
add_bullet('Đây là dự án cá nhân, phục vụ mục đích học tập và thực hành Next.js + Supabase.')
add_bullet('Nền tảng hướng đến cộng đồng sinh viên khu vực Tây Nguyên (chủ yếu Đà Lạt, Bảo Lộc).')
add_bullet('Toàn bộ giao diện và nội dung được viết bằng tiếng Việt.')

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Ngày lập báo cáo: 10/05/2026')
set_font(r, size=11)

# Lưu file
output = 'BaoCao_CampusMarketplace.docx'
doc.save(output)
print(f'[OK] Da tao file: {output}')
